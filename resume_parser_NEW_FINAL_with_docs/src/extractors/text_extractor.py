import os
import logging
import tempfile
import platform
import shutil
import time
import re
import string
import subprocess
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path

# For PDF extraction
from pypdf import PdfReader

# Import PyMuPDF (fitz) for robust PDF extraction
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# For DOCX extraction
import docx2txt
from docx import Document as DocxDocument

# Import olefile for DOC extraction
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

# Import tika for DOC extraction (cross-platform)
try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
    # Initialize tika when first imported
    tika_parser.from_file
except ImportError:
    TIKA_AVAILABLE = False

logger = logging.getLogger(__name__)

class TextExtractor:
    """Extract text from various document formats (PDF, DOCX, DOC, TXT)"""
    
    @staticmethod
    def extract_text(file_path: str, file_type: Optional[str] = None, is_resume: bool = True) -> str:
        """
        Extract text from document based on file type
        
        Args:
            file_path: Path to document file
            file_type: Type of file (pdf, docx, doc, txt) or None to detect from extension
            is_resume: Whether the document is a resume (for specialized cleaning)
            
        Returns:
            Extracted text as string
        """
        # Convert to Path object and resolve to handle any platform differences
        file_path_obj = Path(file_path).resolve()
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path_obj}")
        
        # Determine file type if not provided
        if file_type is None:
            file_extension = file_path_obj.suffix.lower()
            file_type = file_extension.lstrip('.')
        
        file_type = file_type.lower()
        
        # Extract text based on file type
        if file_type == 'pdf':
            text = TextExtractor.extract_from_pdf(str(file_path_obj))
        elif file_type == 'docx':
            text = TextExtractor.extract_from_docx(str(file_path_obj))
        elif file_type == 'doc':
            text = TextExtractor.extract_from_doc(str(file_path_obj))
        elif file_type == 'txt':
            text = TextExtractor.extract_from_txt(str(file_path_obj))
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        # Apply resume-specific cleaning if this is a resume
        if is_resume and text and len(text) > 100:
            # For DOC files (which often have extraction issues), apply more aggressive cleaning
            if file_type == 'doc':
                logger.info(f"Applying resume-specific cleaning to DOC file text: {file_path}")
                return TextExtractor._clean_resume_text(text)
                
        return text
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file with OCR fallback
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        logger.info(f"Extracting text from PDF: {file_path}")
        
        # Try PyMuPDF (fitz) first as it's the most robust method
        if PYMUPDF_AVAILABLE:
            try:
                logger.info("Using PyMuPDF (fitz) to extract text from PDF")
                text = ""
                
                # Open PDF with PyMuPDF
                with fitz.open(file_path) as doc:
                    # Check if document is encrypted/password protected
                    if doc.needs_pass:
                        logger.warning("PDF is password protected, cannot extract text with PyMuPDF")
                    else:
                        # Extract text from each page
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            page_text = page.get_text()
                            text += page_text + "\n\n"
                
                if text.strip() and len(text) > 500:
                    logger.info(f"Extracted {len(text)} chars from PDF using PyMuPDF")
                    return text
                else:
                    logger.warning(f"PyMuPDF extraction produced insufficient text ({len(text.strip())} chars), trying specialized methods")
                    
                    # Try specialized extraction for damaged PDFs
                    special_text = TextExtractor._extract_damaged_pdf_with_pymupdf(file_path)
                    if special_text and len(special_text) > 500:
                        logger.info(f"Specialized extraction successful with {len(special_text)} chars")
                        return special_text
                    logger.warning("Specialized extraction failed, continuing with other methods")
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}, trying alternatives")
        else:
            logger.info("PyMuPDF not available, falling back to other methods")
        
        # Try using PyPDF next
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            if text.strip() and len(text) > 500:
                logger.info(f"Extracted {len(text)} chars from PDF using PyPDF")
                return text
            else:
                logger.warning(f"PyPDF extraction produced insufficient text ({len(text.strip())} chars), trying alternatives")
        except Exception as e:
            logger.warning(f"PyPDF extraction failed: {str(e)}, trying alternatives")
        
        # Try using Tika if available
        if TIKA_AVAILABLE:
            try:
                logger.info("Using Tika to extract text from PDF")
                parsed = tika_parser.from_file(file_path)
                tika_text = parsed.get("content", "")
                
                # Convert to string explicitly if needed - handle case where Tika returns non-string
                if tika_text is not None and not isinstance(tika_text, str):
                    tika_text = str(tika_text)
                
                if tika_text and len(tika_text) > 500:
                    logger.info(f"Extracted {len(tika_text)} chars from PDF using Tika")
                    return tika_text
                else:
                    if tika_text is None:
                        tika_text = ""
                    logger.warning(f"Tika extraction produced insufficient text ({len(tika_text.strip())} chars), trying next method")
            except Exception as e:
                logger.warning(f"Tika extraction failed: {str(e)}, trying next method")
        
        # As a very last resort, try to extract metadata and any embedded text
        try:
            # Get filename as a hint
            import os
            filename = os.path.basename(file_path)
            base_filename = os.path.splitext(filename)[0]
            
            # Try to extract name from filename if it follows Naukri_Name[YY_MM] pattern
            name_match = re.search(r'Naukri_([A-Za-z]+)\[(\d+)y_(\d+)m\]', base_filename)
            if name_match:
                name = name_match.group(1)
                years = name_match.group(2)
                months = name_match.group(3)
                
                # Create minimal text from filename
                minimal_text = f"Name: {name}\nTotal Experience: {years} years {months} months\n"
                
                # Add filename as a hint
                minimal_text += f"File: {filename}\n\n"
                
                logger.info(f"Created minimal text from filename: {minimal_text}")
                return minimal_text
            
            # If we have any text from previous attempts, return it
            if 'text' in locals() and text.strip():
                return text
            elif 'tika_text' in locals() and tika_text and tika_text.strip():
                return tika_text
            else:
                # Return filename at minimum
                return f"File: {filename}\nFailed to extract meaningful text from this PDF."
        except Exception as meta_error:
            logger.error(f"Metadata extraction failed: {str(meta_error)}")
        
        # If all extraction methods failed, return a meaningful error
        logger.error(f"All PDF extraction methods failed for file: {file_path}")
        return f"EXTRACTION_FAILED: Unable to extract text from {os.path.basename(file_path)}"
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            # Use docx2txt for better formatting preservation
            text = docx2txt.process(file_path)
            
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return text
        
        except Exception as e:
            # Fall back to python-docx if docx2txt fails
            try:
                logger.warning(f"docx2txt failed, falling back to python-docx for {file_path}")
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            except Exception as inner_e:
                logger.error(f"Error extracting text from DOCX {file_path}: {str(e)} -> {str(inner_e)}")
                
                # Try Tika as last resort if all else fails
                if TIKA_AVAILABLE:
                    try:
                        logger.info(f"Trying Tika as fallback for DOCX: {file_path}")
                        parsed = tika_parser.from_file(file_path)
                        text = parsed.get("content", "")
                        logger.info(f"Extracted {len(text)} characters from DOCX using Tika: {file_path}")
                        return text
                    except Exception as tika_err:
                        logger.error(f"Tika extraction also failed: {str(tika_err)}")
                
                raise
    
    @staticmethod
    def extract_from_doc(file_path: str) -> str:
        """
        Extract text from DOC file using multiple methods with fallbacks
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Extracted text as string
        """
        logger.info(f"Extracting text from DOC file: {file_path}")
        text = ""
        errors = []
        
        # Method 1: Tika (most reliable for DOC files)
        if TIKA_AVAILABLE:
            try:
                logger.info(f"Using Tika to extract text from DOC: {file_path}")
                parsed = tika_parser.from_file(file_path)
                tika_text = parsed.get("content", "")
                
                if tika_text and len(tika_text) > 100:  # Basic validation
                    logger.info(f"Successfully extracted {len(tika_text)} characters using Tika")
                    text = tika_text
                    # If Tika extraction looks good, return it immediately
                    if len(text.strip()) > 500:  # Substantial content
                        return TextExtractor._clean_doc_text(text)
            except Exception as e:
                errors.append(f"Tika extraction failed: {str(e)}")
                logger.warning(f"Tika extraction failed: {str(e)}")
        else:
            errors.append("Tika not available")
        
        # Method 2: Use olefile (if available)
        if not text and OLEFILE_AVAILABLE:
            try:
                import olefile
                logger.info(f"Using olefile to extract text from DOC: {file_path}")
                
                if olefile.isOleFile(file_path):
                    ole_text = ""
                    # Extract text from the OLE file
                    with olefile.OleFile(file_path) as ole:
                        # Try to read WordDocument stream
                        if ole.exists('WordDocument'):
                            try:
                                word_data = ole.openstream('WordDocument').read()
                                logger.info(f"Read WordDocument stream: {len(word_data)} bytes")
                                
                                # Try to extract text from WordDocument stream
                                content = word_data.replace(b'\x00', b' ').decode('utf-8', errors='ignore')
                                content = ''.join(char if char.isprintable() or char in ['\n', '\r', '\t'] else ' ' for char in content)
                                content = re.sub(r'\s+', ' ', content).strip()
                                
                                if len(content) > 100:
                                    ole_text += content + "\n"
                            except Exception as e:
                                logger.warning(f"WordDocument stream extraction failed: {str(e)}")
                        
                        # Extract text from additional streams that might contain text
                        for stream_name in ['1Table', 'Table', '0Table', 'Data']:
                            if ole.exists(stream_name):
                                try:
                                    stream_data = ole.openstream(stream_name).read()
                                    try:
                                        stream_content = stream_data.replace(b'\x00', b' ').decode('utf-8', errors='ignore')
                                        stream_content = ''.join(char for char in stream_content if char.isprintable() or char in ['\n', '\r', '\t'])
                                        
                                        # Add only if substantial content
                                        if len(stream_content.strip()) > 50:
                                            ole_text += stream_content + "\n"
                                    except Exception as e:
                                        logger.warning(f"Error extracting text from {stream_name}: {str(e)}")
                                except Exception as e:
                                    logger.warning(f"Error reading stream {stream_name}: {str(e)}")
                        
                        # Get summary information (metadata)
                        if ole.exists('\x05SummaryInformation'):
                            try:
                                summary = ole.getproperties('\x05SummaryInformation')
                                if summary:
                                    for key, value in summary.items():
                                        if isinstance(value, str) and len(value) > 10:
                                            if key == 2:  # Title
                                                ole_text = f"Title: {value}\n{ole_text}"
                                            elif key == 3:  # Subject
                                                ole_text = f"{ole_text}\nSubject: {value}"
                                            elif key == 4:  # Author
                                                ole_text = f"{ole_text}\nAuthor: {value}"
                            except Exception as e:
                                logger.warning(f"Error extracting summary information: {str(e)}")
                    
                    if ole_text and len(ole_text) > 100:
                        if not text or len(ole_text) > len(text):
                            text = ole_text
                            logger.info(f"Successfully extracted {len(text)} characters using olefile")
            except Exception as e:
                errors.append(f"OleFile extraction failed: {str(e)}")
                logger.warning(f"OleFile extraction failed: {str(e)}")
        
        # Method 3: Extract filename as a hint for the resume
        filename = os.path.basename(file_path)
        filename_without_ext = os.path.splitext(filename)[0]
        
        # Try to extract name from filename if it follows Naukri_Name[YY_MM] pattern
        name_match = re.search(r'Naukri_([A-Za-z]+)\[(\d+)y_(\d+)m\]', filename_without_ext)
        if name_match:
            name = name_match.group(1)
            years = name_match.group(2)
            months = name_match.group(3)
            filename_text = f"Name: {name}\nExperience: {years} years {months} months\n\n"
            text = filename_text + text
            logger.info(f"Added name and experience from filename: {name}, {years}y {months}m")
        
        # If text is still empty or very short, add the filename as a hint
        if len(text.strip()) < 100:
            text = f"Filename: {filename_without_ext}\n\n{text}"
            logger.warning(f"Extracted minimal text, added filename as hint: {filename_without_ext}")
        
        # Apply cleaning to improve DOC text
        cleaned_text = TextExtractor._clean_doc_text(text)
        
        if len(cleaned_text.strip()) < 100:
            logger.warning(f"DOC extraction produced minimal text ({len(cleaned_text)} chars) after cleaning")
        else:
            logger.info(f"Final DOC extraction: {len(cleaned_text)} characters")
            
        return cleaned_text
    
    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """
        Extract text from plaintext file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            
            logger.info(f"Extracted {len(text)} characters from TXT file: {file_path}")
            return text
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text = f.read()
                    logger.info(f"Extracted {len(text)} characters from TXT file using {encoding} encoding: {file_path}")
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use binary mode as last resort
            with open(file_path, 'rb') as f:
                binary_data = f.read()
                text = binary_data.decode('utf-8', errors='replace')
            
            logger.warning(f"Used fallback binary mode for TXT file: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT file {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_metadata(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to document file
            file_type: Type of file (pdf, docx, doc, txt) or None to detect from extension
            
        Returns:
            Dictionary of metadata
        """
        file_path_obj = Path(file_path).resolve()
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path_obj}")
        
        # Determine file type if not provided
        if file_type is None:
            file_extension = file_path_obj.suffix.lower()
            file_type = file_extension.lstrip('.')
        
        file_type = file_type.lower()
        file_size = file_path_obj.stat().st_size
        file_name = file_path_obj.name
        
        metadata = {
            'file_name': file_name,
            'file_type': file_type,
            'file_size': file_size,
            'creation_date': None,
            'modification_date': None,
            'author': None,
            'title': None
        }
        
        # Try to get metadata using Tika for any format
        if TIKA_AVAILABLE:
            try:
                logger.debug(f"Extracting metadata with Tika for {file_path}")
                parsed = tika_parser.from_file(file_path)
                tika_meta = parsed.get("metadata", {})
                
                # Map Tika metadata fields to our fields
                if "Author" in tika_meta:
                    metadata['author'] = tika_meta["Author"]
                elif "creator" in tika_meta:
                    metadata['author'] = tika_meta["creator"]
                    
                if "title" in tika_meta:
                    metadata['title'] = tika_meta["title"]
                
                if "Creation-Date" in tika_meta:
                    metadata['creation_date'] = tika_meta["Creation-Date"]
                
                if "Last-Modified" in tika_meta:
                    metadata['modification_date'] = tika_meta["Last-Modified"]
                
                return metadata
            except Exception as e:
                logger.warning(f"Tika metadata extraction failed: {str(e)}")
        
        # Extract specific metadata based on file type using other methods
        try:
            if file_type == 'pdf':
                pdf_reader = PdfReader(str(file_path_obj))
                info = pdf_reader.metadata
                if info:
                    metadata['author'] = info.author
                    metadata['title'] = info.title
                    metadata['creation_date'] = info.creation_date
                    metadata['modification_date'] = info.modification_date
            
            elif file_type == 'docx':
                doc = DocxDocument(str(file_path_obj))
                core_props = doc.core_properties
                metadata['author'] = core_props.author
                metadata['title'] = core_props.title
                metadata['creation_date'] = core_props.created
                metadata['modification_date'] = core_props.modified
            
            elif file_type == 'txt':
                # For text files, just get file stats
                stat_info = file_path_obj.stat()
                import datetime
                metadata['creation_date'] = datetime.datetime.fromtimestamp(stat_info.st_ctime)
                metadata['modification_date'] = datetime.datetime.fromtimestamp(stat_info.st_mtime)
        
        except Exception as e:
            logger.warning(f"Error extracting metadata from {file_path}: {str(e)}")
        
        return metadata 
    
    @staticmethod
    def _clean_doc_text(text: str) -> str:
        """
        Special cleaning for DOC text which often has extraction artifacts
        
        Args:
            text: Raw text from DOC file
            
        Returns:
            Cleaned text
        """
        # Remove non-printable characters except newlines and tabs
        text = ''.join(char if char.isprintable() or char in ['\n', '\r', '\t'] else ' ' for char in text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove empty lines at start
        text = re.sub(r'^\s*\n', '', text)
        
        # Remove common DOC extraction artifacts
        text = re.sub(r'Evaluation Warning.*?document\..*?\n', '', text, flags=re.IGNORECASE|re.DOTALL)
        text = re.sub(r'Microsoft Word.*?Document.*?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Normal\.dot.*?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'MSWordDoc.*?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'MSWorksWPDoc.*?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'objCh.*?\n', '', text, flags=re.IGNORECASE)
        
        # Clean up extra null characters and their remnants
        text = text.replace('\x00', ' ')
        text = re.sub(r'\s+\n', '\n', text)
        text = re.sub(r'\n\s+', '\n', text)
        
        return text.strip()
    
    @staticmethod
    def _clean_resume_text(text: str) -> str:
        """
        Clean up extracted resume text to improve parsing quality
        
        Args:
            text: Raw extracted text from a resume
            
        Returns:
            Cleaned resume text
        """
        # Replace multiple spaces with single spaces
        text = re.sub(r' {2,}', ' ', text)
        
        # Replace multiple newlines with double newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove spaces at end of lines
        text = re.sub(r'[^\S\n]+\n', '\n', text)
        
        # Remove spaces at beginning of lines
        text = re.sub(r'\n[^\S\n]+', '\n', text)
        
        return text
    
    @staticmethod
    def _extract_damaged_pdf_with_pymupdf(file_path: str) -> str:
        """
        Specialized extraction method for damaged or problematic PDFs
        Uses PyMuPDF's advanced capabilities to extract both text and images
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string, or empty string if extraction fails
        """
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not available for damaged PDF extraction")
            return ""
        
        try:
            logger.info(f"Attempting specialized extraction for damaged PDF: {file_path}")
            result_text = ""
            
            with fitz.open(file_path) as doc:
                # Try text extraction with different strategies
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Try with different extraction flags
                    flags = fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE
                    page_text = page.get_text("text", flags=flags)
                    
                    # If standard extraction failed, try extracting as HTML
                    if not page_text.strip():
                        logger.debug("Standard text extraction failed, trying HTML extraction")
                        html_text = page.get_text("html")
                        
                        # Extract readable text from HTML
                        if html_text:
                            # Use regex to extract text from HTML tags
                            text_parts = re.findall(r'>([^<]+)<', html_text)
                            page_text = ' '.join(text_parts)
                    
                    # If text extraction still failed, try raw extraction
                    if not page_text.strip():
                        logger.debug("HTML extraction failed, trying raw extraction")
                        raw_dict = page.get_text("dict")
                        if "blocks" in raw_dict:
                            for block in raw_dict["blocks"]:
                                if "lines" in block:
                                    for line in block["lines"]:
                                        if "spans" in line:
                                            for span in line["spans"]:
                                                if "text" in span:
                                                    page_text += span["text"] + " "
                    
                    result_text += page_text + "\n\n"
            
            if result_text.strip():
                logger.info(f"Successfully extracted {len(result_text)} chars with specialized method")
                return result_text
            else:
                logger.warning("Specialized text extraction failed to get usable content")
                return ""
            
        except Exception as e:
            logger.error(f"Error in specialized PDF extraction: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_pdf_raw_streams(pdf_path: str) -> str:
        """
        Extract any readable text directly from PDF binary data
        Used as a last resort for severely corrupted PDFs
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text from binary data
        """
        logger.info(f"Attempting to extract raw PDF streams from: {pdf_path}")
        try:
            with open(pdf_path, 'rb') as f:
                data = f.read()
            
            # Look for content streams
            text_chunks = []
            
            # Find text objects in PDF streams (between BT and ET markers)
            for match in re.finditer(b'BT(.*?)ET', data, re.DOTALL):
                chunk = match.group(1)
                # Extract text strings (between parentheses)
                for text_match in re.finditer(b'\\((.*?)\\)', chunk):
                    try:
                        text = text_match.group(1).decode('utf-8', errors='replace')
                        if len(text) > 3:  # Filter out short garbage
                            text_chunks.append(text)
                    except:
                        pass
            
            # Find dictionary objects that might contain text
            for match in re.finditer(b'/Contents\\s*\\((.*?)\\)', data, re.DOTALL):
                try:
                    text = match.group(1).decode('utf-8', errors='replace')
                    if len(text) > 5:  # Filter out short garbage
                        text_chunks.append(text)
                except:
                    pass
                
            # Find any strings in the PDF that look like they might be content
            for match in re.finditer(b'/Text\\s*\\((.*?)\\)', data, re.DOTALL):
                try:
                    text = match.group(1).decode('utf-8', errors='replace')
                    if len(text) > 5:  # Filter out short garbage
                        text_chunks.append(text)
                except:
                    pass
            
            # Extract anything that looks like it might be regular text
            # Look for decent-sized runs of printable ASCII characters
            text_pattern = re.compile(b'[A-Za-z0-9\\s.,;:\\-\'\"]{20,}')
            for match in text_pattern.finditer(data):
                try:
                    text = match.group(0).decode('utf-8', errors='replace')
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 20:  # Filter out short garbage
                        text_chunks.append(text)
                except:
                    pass
            
            extracted_text = "\n".join(text_chunks)
            logger.info(f"Extracted {len(text_chunks)} text chunks ({len(extracted_text)} chars) from raw PDF data")
            
            return extracted_text
        except Exception as e:
            logger.error(f"Raw PDF stream extraction failed: {str(e)}")
            return "" 